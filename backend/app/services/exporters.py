"""Local export of generated test cases to CSV / DOCX."""

from __future__ import annotations

import csv
from datetime import date
from pathlib import Path

from app.schemas.test_cases import TestCase

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[misc, assignment]


def default_export_dir(ticket_key: str) -> Path:
    return Path("output") / "test_cases" / ticket_key


def filename_for(ticket_key: str, fmt: str, day: date | None = None) -> str:
    stamp = (day or date.today()).isoformat()
    return f"{ticket_key}_{stamp}.{fmt}"


def write_csv(path: Path, cases: list[TestCase]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow(["id", "title", "type", "steps", "expected_result"])
        for case in cases:
            writer.writerow(
                [
                    case.id,
                    case.title,
                    case.type,
                    " | ".join(case.steps),
                    case.expected_result,
                ]
            )
    return path


def write_docx(path: Path, ticket_key: str, cases: list[TestCase]) -> Path:
    if Document is None:
        raise RuntimeError("python-docx is required for DOCX export. Install backend requirements.")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(f"Test cases — {ticket_key}", level=1)
    for case in cases:
        doc.add_heading(f"{case.id} · {case.title}", level=2)
        doc.add_paragraph(f"Type: {case.type}")
        doc.add_paragraph("Steps:")
        for index, step in enumerate(case.steps, start=1):
            doc.add_paragraph(f"{index}. {step}", style="List Number")
        doc.add_paragraph(f"Expected result: {case.expected_result}")
    doc.save(path)
    return path


def export_cases(
    folder: Path,
    ticket_key: str,
    cases: list[TestCase],
    formats: list[str],
) -> dict[str, str]:
    saved: dict[str, str] = {}
    for fmt in formats:
        path = folder / filename_for(ticket_key, fmt)
        if fmt == "csv":
            write_csv(path, cases)
        elif fmt == "docx":
            write_docx(path, ticket_key, cases)
        else:
            continue
        saved[fmt] = str(path).replace("\\", "/")
    return saved


def read_export_bytes(path: Path) -> bytes:
    return path.read_bytes()
